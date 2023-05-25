from PyQt5.QtWidgets import *
from PyQt5.QtWebEngineWidgets import QWebEngineSettings, QWebEngineView
from PyQt5 import QAxContainer
from PyQt5.QtCore import Qt
from PyQt5.QtWidgets import QWidget, QVBoxLayout, QPushButton, QLineEdit, QApplication,  QProgressBar, QComboBox, QMessageBox
from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtCore import QThread, pyqtSignal
from PyQt5.QtCore import QSize
import PyPDF2
import sys
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert
import time
import glob
from datetime import date
from docxtpl import DocxTemplate
from cap_image import MainWindow

today = date.today()

class Window(QMainWindow):
    def __init__(self):
        super().__init__()
        # set the title of main window
        self.setWindowTitle('Barangay Calamias Document System')
        self.num = 0

        # Menu bars
        
        # set the size of window
        self.Width = 800
        self.height = int(0.618 * self.Width)
        self.resize(self.Width, self.height)

        # add all widgets
        self.btn_1 = QPushButton('Barangay Clearance', self)
        self.btn_1.setStyleSheet('QPushButton{background-color: rgb(0, 170, 0); color: rgb(255, 255, 255)}')
        self.btn_2 = QPushButton('Barangay Indigency', self)
        self.btn_2.setStyleSheet('QPushButton{background-color: rgb(0, 170, 0); color: rgb(255, 255, 255)}')
        self.btn_3 = QPushButton('Barangay Certificate', self)
        self.btn_3.setStyleSheet('QPushButton{background-color: rgb(0, 170, 0);color: rgb(255, 255, 255)}')
        self.btn_4 = QPushButton('Barangay Business Clearance', self)
        self.btn_4.setStyleSheet('QPushButton{background-color: rgb(0, 170, 0); color: rgb(255, 255, 255)}')
        self.btn_5 = QPushButton('Certificate of Guardianship', self)
        self.btn_5.setStyleSheet('QPushButton{background-color: rgb(0, 170, 0); color: rgb(255, 255, 255)}')
        self.btn_6 = QPushButton('Putol Puno', self)
        self.btn_6.setStyleSheet('QPushButton{background-color: rgb(0, 170, 0); color: rgb(255, 255, 255)}')
        self.btn_7 = QPushButton('About', self)
        self.btn_7.setStyleSheet('QPushButton{background-color: rgb(0, 170, 0); color: rgb(255, 255, 255)}')

        self.btn_1.clicked.connect(self.button1)
        self.btn_2.clicked.connect(self.button2)
        self.btn_3.clicked.connect(self.button3)
        self.btn_4.clicked.connect(self.button4)
        self.btn_5.clicked.connect(self.button5)
        self.btn_6.clicked.connect(self.button6)
        self.btn_7.clicked.connect(self.button7)

        # add tabs
        self.tab1 = self.ui1()
        self.tab2 = self.ui2()
        self.tab3 = self.ui3()
        self.tab4 = self.ui4()
        self.tab5 = self.ui5()
        self.tab6 = self.ui6()
        self.tab7 = self.ui7()

        self.initUI()

    def initUI(self):
        left_layout = QVBoxLayout()

        # Logo
        label =  QtWidgets.QLabel()
        label.setMaximumSize(QtCore.QSize(150, 150))
        label.setText("")
        p = sys.path.append("C:/resources")
        label.setPixmap(QtGui.QPixmap(r'C:\Users\Angeline\Desktop\OJT-SYSTEM\calamias-system\img\logo.png'))
        label.setScaledContents(True)
        label.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        label.setIndent(-1)
        label.setObjectName('logo')
         
        left_layout.addWidget(label)
        left_layout.addWidget(self.btn_1)
        left_layout.addWidget(self.btn_2)
        left_layout.addWidget(self.btn_3)
        left_layout.addWidget(self.btn_4)
        left_layout.addWidget(self.btn_5)
        left_layout.addWidget(self.btn_6)
        left_layout.addWidget(self.btn_7)
        left_layout.addWidget(QLabel('KLL BSCS OJT 2023'))
        left_layout.addStretch(5)
        left_layout.setSpacing(20)
        label.setStyleSheet('QLabel{background-color:#FFFFFF}')
        left_widget = QWidget()
        left_widget.setLayout(left_layout)

        self.right_widget = QTabWidget()
        self.right_widget.tabBar().setObjectName("mainTab")
        self.right_widget.addTab(self.tab1, '')
        self.right_widget.addTab(self.tab2, '')
        self.right_widget.addTab(self.tab3, '')
        self.right_widget.addTab(self.tab4, '')
        self.right_widget.addTab(self.tab5, '')
        self.right_widget.addTab(self.tab6, '')
        self.right_widget.addTab(self.tab7, '')

        self.right_widget.setCurrentIndex(0)
        self.right_widget.setStyleSheet('''QTabBar::tab{width: 0; \
            height: 0; margin: 0; padding: 0; border: none;}''')

        main_layout = QHBoxLayout()
        main_layout.addWidget(left_widget)
        # frame = QFrame(self)
        # frame.setFrameShape(QFrame.StyledPanel)
        # frame.setGeometry(20, 20, 360, 260)
        # frame.setStyleSheet("background-color: #f0f0f0;")
        # main_layout.addWidget(frame)
        main_layout.addWidget(self.right_widget)
        main_layout.setStretch(0, 40)
        main_layout.setStretch(1, 300)  
        main_widget = QWidget()
        main_widget.setLayout(main_layout)
        self.setCentralWidget(main_widget)

    # ----------------- 
    # buttons

    def button1(self):
        self.right_widget.setCurrentIndex(0)

    def button2(self):
        self.right_widget.setCurrentIndex(1)

    def button3(self):
        self.right_widget.setCurrentIndex(2)

    def button4(self):
        self.right_widget.setCurrentIndex(3)

    def button5(self):
        self.right_widget.setCurrentIndex(4)

    def button6(self):
        self.right_widget.setCurrentIndex(5)

    def button7(self):
        self.right_widget.setCurrentIndex(6)


    # ----------------- 
    # pages
    
    def ui1(self):
        WebBrowser = QVBoxLayout()
        container = QAxContainer.QAxWidget()
        # container.setStyleSheet('''width: 0; \
        #     height: 0; margin: 0; padding: 0; border: none; display: inline-block;''')
        container.setFocusPolicy(Qt.StrongFocus)
        container.setControl("{8856F961-340A-11D0-A96B-00C04FD705A2}")

        box = QGroupBox("Barangay Clearance")
        group_box_layout = QFormLayout()
        image_capture = QPushButton('Capture image')
        image_capture.setStyleSheet('QPushButton{background-color:  rgb(255, 0, 0); color: rgb(255, 255, 255)}')
        text_box = QLineEdit()
        text_box.move(20, 80)
        text_box.resize(280, 40)
        btn = QPushButton('Proceed')
        btn.setStyleSheet('QPushButton{background-color: rgb(0, 170, 0); color: rgb(255, 255, 255)}')

        WebBrowser.addWidget(box)
        WebBrowser.addWidget(container)
        ctrl_no = 0

        def print_clearance(self):
            doc = Document(r'brgy-certificate\brgy-docx-file\brgy-clearance.docx')
            ts = 0
            found = None
            loc = r"C:\Users\Angeline\Desktop\OJT-SYSTEM\calamias-system\capture-image-location\*"

            for file_name in glob.glob(loc):
                fts = os.path.getmtime(file_name)
                if fts > ts:
                    ts = fts
                    found = file_name

            text_to_find = 'replace'
            for paragraph in doc.paragraphs:
                if text_to_find in paragraph.text:
                    run = paragraph.runs[3]  # Get the first run in the paragraph
                    run.clear()  # Clear the existing text

                    # Insert the image
                    run.add_picture(found, width=Inches(1))

            
            Dictionary = {
                            "John Mark S. Salcedo ": text_box.text().upper() + ' ',
                            "May": today.strftime('%B'), 
                            "15": today.strftime('%d'),
                            "2023": today.strftime('%Y'),
                            "Control No": "Control No"
                        }

            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(12)
            font.color.rgb = RGBColor(59, 56, 56)

            for i in Dictionary:
                for p in doc.paragraphs:
                    if p.text.find(i) >= 0:
                        p.text = p.text.replace(i,Dictionary[i])

            #save changed document

            try:
                # time.sleep(3)
                f = r"C:\Users\Angeline\Desktop\OJT-SYSTEM\calamias-system\brgy-certificate\brgy-docx-file\clearance-files\clearance_results.docx"
                file = doc.save(f)
          
                convert(f, r'C:\Users\Angeline\Desktop\OJT-SYSTEM\calamias-system\brgy-certificate\brgy-docx-file\clearance-files')
                print('done.....')

                df = r"C:\Users\Angeline\Desktop\OJT-SYSTEM\calamias-system\brgy-certificate\brgy-docx-file\clearance-files\clearance_results.pdf"
                container.dynamicCall('Navigate(const QString&)', df)

            except PermissionError:
                QMessageBox.about(self, 'error', 'opps you need to close the file first')

        def capture_image():
            cap = MainWindow()
            cap.show()
            print('image done...')
            
        image_capture.clicked.connect(capture_image)
               
        btn.clicked.connect(print_clearance)
        group_box_layout.addRow(image_capture)
        group_box_layout.addRow(QLabel('Name: '), text_box)
        group_box_layout.addRow(btn)
        box.setLayout(group_box_layout)
        
        main = QWidget()
        main.setLayout(WebBrowser)

        return main

    def ui2(self):
        WebBrowser = QVBoxLayout()
        
        # time.sleep(5)
        container = QAxContainer.QAxWidget()
        container.setControl("{8856F961-340A-11D0-A96B-00C04FD705A2}")
        
        box = QGroupBox("Barangay Indigency")
        group_box_layout = QFormLayout()
        text_box = QLineEdit()
        text_box.move(20, 80)
        text_box.resize(280, 40)
        btn = QPushButton('Proceed')
        btn.setStyleSheet('QPushButton{background-color: rgb(0, 170, 0); color: rgb(255, 255, 255)}')

        # make indigency certificate as pdf file 
        def print_indigency():
            doc = Document('brgy-certificate/brgy-docx-file/brgy-indigency.docx')
            Dictionary = {
                            "John Mark S. Salcedo ": text_box.text() + ' '.upper(),
                            "May": today.strftime('%B'), 
                            "15": today.strftime('%d'),
                            "2023": today.strftime('%Y')
                        }
            
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(11)
            font.color.rgb = RGBColor(59, 56, 56)

            for i in Dictionary:
                for p in doc.paragraphs:
                    if p.text.find(i) >= 0:
                        p.text = p.text.replace(i,Dictionary[i])

            #save changed document
            
            try:
                # time.sleep(3)
                f = r"C:\Users\Angeline\Desktop\OJT-SYSTEM\calamias-system\brgy-certificate\brgy-docx-file\converted-files\indigency_results.docx"
                file = doc.save(f)
          
                convert(f, r'C:\Users\Angeline\Desktop\OJT-SYSTEM\calamias-system\brgy-certificate\brgy-docx-file\converted-files')
                print('done.....')
                df = r"C:\Users\Angeline\Desktop\OJT-SYSTEM\calamias-system\brgy-certificate\brgy-docx-file\converted-files"
                container.dynamicCall('Navigate(const QString&)', df)

            except PermissionError:
                QMessageBox.about(self, 'error', 'opps you need to close the file first')
            
        btn.clicked.connect(print_indigency)
        
        group_box_layout.addRow(QLabel('Name: '), text_box)
        group_box_layout.addRow(btn)
        box.setLayout(group_box_layout)

        WebBrowser.addWidget(box)
        WebBrowser.addWidget(container)
        main = QWidget()
        main.setLayout(WebBrowser)

        return main

    def ui3(self):
        WebBrowser = QVBoxLayout()
        container = QAxContainer.QAxWidget()
        # container.setStyleSheet('''width: 0; \
        #     height: 0; margin: 0; padding: 0; border: none; display: inline-block;''')
        container.setFocusPolicy(Qt.StrongFocus)
        container.setControl("{8856F961-340A-11D0-A96B-00C04FD705A2}")
        box = QGroupBox("Barangay Certificate")
        group_box_layout = QFormLayout()
        image_capture = QPushButton('Capture image')
        image_capture.setStyleSheet('QPushButton{background-color:  rgb(255, 0, 0); color: rgb(255, 255, 255)}')
        lname = QLineEdit()
        fname = QLineEdit()
        mname = QLineEdit()
        since = QLineEdit()
        nick_name = QLineEdit()
        bdate = QDateEdit()
        birth_place = QLineEdit()
        age = QLineEdit()
        status = QComboBox()
        status.addItems(['single', 'married', 'widow', 'legaly separated'])
        gender = QComboBox()
        gender.addItems(['male', 'female', 'other'])
        purpose = QLineEdit()
        btn = QPushButton('Proceed')
        btn.setStyleSheet('QPushButton{background-color: rgb(0, 170, 0); color: rgb(255, 255, 255)}')

        # make indigency certificate as pdf file 
        def print_certificate():
            doc = Document('brgy-certificate/brgy-docx-file/brgy-certificate.docx')
            
            tables = doc.tables
            ts = 0
            found = None
            loc = r"C:\Users\Angeline\Desktop\OJT-SYSTEM\calamias-system\capture-image-location\*"
            for file_name in glob.glob(loc):
                fts = os.path.getmtime(file_name)
                if fts > ts:
                    ts = fts
                    found = file_name

            text_to_find = 'replace'
            for paragraph in doc.paragraphs:
                if text_to_find in paragraph.text:
                    run = paragraph.runs[4]  # Get the first run in the paragraph
                    run.clear()  # Clear the existing text

                    # Insert the image
                    run.add_picture(found, width=Inches(1))   

            birth_date = str(bdate.date().toPyDate()).replace('-', ' ')
            val = list(birth_date.split(" "))
            arr = []

            for x in val:
                arr.append(int(x))

            format_datetime = date(arr[0], arr[1], arr[2])

            Dictionary = {
                            "SALCEDO": ' ' +lname.text().upper(),
                            "JOHN": fname.text().upper(),
                            "SATURNO": mname.text().upper(),
                            "BROWNY": nick_name.text().upper(),
                            "May": today.strftime('%B'), 
                            "15": today.strftime('%d'),
                            "2023": today.strftime('%Y'),
                            "1999": since.text(),
                            "APRIL 09,1982": format_datetime.strftime('%b %d,%Y'),
                            "QUIAPO": birth_place.text().upper(),
                            "51": age.text(),
                            "SINGLE": status.currentText().upper(),
                            "MALE": gender.currentText().upper(),
                            "As temporary ID/proof of identification": purpose.text().upper(),
                            "FEBRUARY 16, 20223": today.strftime('%B %d, %Y')
                        }

            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(12)
            font.color.rgb = RGBColor(59, 56, 56)
            
            for i in Dictionary:
                for p in doc.paragraphs:
                    if p.text.find(i) >= 0:
                        p.text = p.text.replace(i,Dictionary[i])
            
            try:
                # time.sleep(3)
                f = r"C:\Users\Angeline\Desktop\OJT-SYSTEM\calamias-system\brgy-certificate\brgy-docx-file\certificate-files\certificate_results.docx"
                file = doc.save(f)
          
                convert(f, r'C:\Users\Angeline\Desktop\OJT-SYSTEM\calamias-system\brgy-certificate\brgy-docx-file\certificate-files')
                print('done.....')
                df = "file:///C:/Users/Angeline/Desktop/OJT-SYSTEM/calamias-system/brgy-certificate/brgy-docx-file/certificate-files"
                container.dynamicCall('Navigate(const QString&)', df)

            except PermissionError:
                QMessageBox.about(self, 'error', 'opps you need to close the file first')
      
        def capture_image():
            cap = MainWindow()
            cap.show()
            print('image done...')
            
        btn.clicked.connect(print_certificate)
        image_capture.clicked.connect(capture_image)
        
        group_box_layout.addRow(image_capture)
        group_box_layout.addRow(QLabel('Last Name: '), lname)
        group_box_layout.addRow(QLabel('First Name: '), fname)
        group_box_layout.addRow(QLabel('Middle Name: '), mname)
        group_box_layout.addRow(QLabel('Nick Name: '), nick_name)
        group_box_layout.addRow(QLabel('Since: '), since)
        group_box_layout.addRow(QLabel('Birth Date: '), bdate)
        group_box_layout.addRow(QLabel('Birth Place: '), birth_place)
        group_box_layout.addRow(QLabel('Age: '), age)
        group_box_layout.addRow(QLabel('Status: '), status)
        group_box_layout.addRow(QLabel('Gender: '), gender)
        group_box_layout.addRow(QLabel('Purpose: '), purpose)
        group_box_layout.addRow(btn)
        box.setLayout(group_box_layout)

        WebBrowser.addWidget(box)
        WebBrowser.addWidget(container)
        main = QWidget()
        main.setLayout(WebBrowser)

        return main

    def ui4(self):
        WebBrowser = QVBoxLayout()
        container = QAxContainer.QAxWidget()
        # container.setStyleSheet('''width: 0; \
        #     height: 0; margin: 0; padding: 0; border: none; display: inline-block;''')
        container.setFocusPolicy(Qt.StrongFocus)
        container.setControl("{8856F961-340A-11D0-A96B-00C04FD705A2}")
        # WebBrowser.addWidget(QCalendarWidget())box = QGroupBox("Barangay Indigency")
        box = QGroupBox("Barangay Business Clearance")
        group_box_layout = QFormLayout()
        name = QLineEdit()
        businessname = QLineEdit()
        location = QComboBox()
        location.addItems(['Purok 1', 'Purok 2', 'Purok 3', 'Purok 4', 'Purok 5', 'Purok 6'])
        btn = QPushButton('Proceed')
        btn.setStyleSheet('QPushButton{background-color: rgb(0, 170, 0); color: rgb(255, 255, 255)}')

        # make indigency certificate as pdf file 
        def print_businesscleareance():
            doc = Document('brgy-certificate/brgy-docx-file/brgy-businessclearance.docx')

            Dictionary = {
                            "James Ried": name.text().upper() + ' ',
                            "JAMES LID ": businessname.text().upper() + ' ',
                            "PUROKYA NI JAMES RIED": location.currentText().upper() + ' ' + 'Calamias, Ibaan, Batangas'.upper(),
                            "MAY": today.strftime('%B'), 
                            "11": today.strftime('%d'),
                            "2023": today.strftime('%Y')
                        }

            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(12)
            font.color.rgb = RGBColor(59, 56, 56)

            for i in Dictionary:
                for p in doc.paragraphs:
                    if p.text.find(i) >= 0:
                        p.text = p.text.replace(i,Dictionary[i])

            #save changed document
            
            try:
                # time.sleep(3)
                f = r"C:\Users\Angeline\Desktop\OJT-SYSTEM\calamias-system\brgy-certificate\brgy-docx-file\businessclearances-files\businessclearance_results.docx"
                file = doc.save(f)
          
                convert(f, r'C:\Users\Angeline\Desktop\OJT-SYSTEM\calamias-system\brgy-certificate\brgy-docx-file\businessclearances-files')
                print('done.....')
                df = "file:///C:/Users/Angeline/Desktop/OJT-SYSTEM/calamias-system/brgy-certificate/brgy-docx-file/businessclearances-files"
                container.dynamicCall('Navigate(const QString&)', df)

            except PermissionError:
                QMessageBox.about(self, 'error', 'opps you need to close the file first')
            
        btn.clicked.connect(print_businesscleareance)
        
        group_box_layout.addRow(QLabel('Proprietor: '), name)
        group_box_layout.addRow(QLabel('Business/Trade Name: '), businessname)
        group_box_layout.addRow(QLabel('Location: '), location)
        group_box_layout.addRow(btn)
        box.setLayout(group_box_layout)

        WebBrowser.addWidget(box)
        WebBrowser.addWidget(container)
        main = QWidget()
        main.setLayout(WebBrowser)

        return main

    def ui5(self):
        WebBrowser = QVBoxLayout()
        container = QAxContainer.QAxWidget()
        # container.setStyleSheet('''width: 0; \
        #     height: 0; margin: 0; padding: 0; border: none; display: inline-block;''')
        container.setFocusPolicy(Qt.StrongFocus)
        container.setControl("{8856F961-340A-11D0-A96B-00C04FD705A2}")
        # WebBrowser.addWidget(QCalendarWidget())box = QGroupBox("Barangay Indigency")
        box = QGroupBox("Barangay Certificate of Guardianship")
        group_box_layout = QFormLayout()
        guardian = QLineEdit()
        ward = QLineEdit()
        btn = QPushButton('Proceed')
        btn.setStyleSheet('QPushButton{background-color: rgb(0, 170, 0); color: rgb(255, 255, 255)}')

        # make indigency certificate as pdf file 
        def print_guardianship():
            doc = Document('brgy-certificate/brgy-docx-file/brgy-guardianship.docx')

            Dictionary = {
                            "Francis Indino": guardian.text().upper() + ' ',
                            "John Mark Salcedo": ward.text().upper() + ' ',
                            "May": today.strftime('%B'), 
                            "15": today.strftime('%d'),
                            "2023": today.strftime('%Y')
                        }

            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(12)
            font.color.rgb = RGBColor(59, 56, 56)

            for i in Dictionary:
                for p in doc.paragraphs:
                    if p.text.find(i) >= 0:
                        p.text = p.text.replace(i,Dictionary[i])

            #save changed document
            
            try:
                # time.sleep(3)
                f = r"C:\Users\Angeline\Desktop\OJT-SYSTEM\calamias-system\brgy-certificate\brgy-docx-file\guardianship-files\guardianship_results.docx"
                file = doc.save(f)
          
                convert(f, r'C:\Users\Angeline\Desktop\OJT-SYSTEM\calamias-system\brgy-certificate\brgy-docx-file\guardianship-files')
                print('done.....')
                df = "file:///C:/Users/Angeline/Desktop/OJT-SYSTEM/calamias-system/brgy-certificate/brgy-docx-file/guardianship-files"
                container.dynamicCall('Navigate(const QString&)', df)

            except PermissionError:
                 QMessageBox.about(self, 'error', 'opps you need to close the file first')
               
            
        btn.clicked.connect(print_guardianship)
        
        group_box_layout.addRow(QLabel('Guardian Name: '), guardian)
        group_box_layout.addRow(QLabel('Ward Name: '), ward)
        group_box_layout.addRow(btn)
        box.setLayout(group_box_layout)

        WebBrowser.addWidget(box)
        WebBrowser.addWidget(container)
        main = QWidget()
        main.setLayout(WebBrowser)

        return main
        
    def ui6(self):
        WebBrowser = QVBoxLayout()
        container = QAxContainer.QAxWidget()
        # container.setStyleSheet('''width: 0; \
        #     height: 0; margin: 0; padding: 0; border: none; display: inline-block;''')
        container.setFocusPolicy(Qt.StrongFocus)
        container.setControl("{8856F961-340A-11D0-A96B-00C04FD705A2}")
        # WebBrowser.addWidget(QCalendarWidget())box = QGroupBox("Barangay Indigency")
        box = QGroupBox("Putol Puno")
        group_box_layout = QFormLayout()
        name = QLineEdit()
        number_of_trees = QLineEdit()
        tree = QLineEdit()
        purpose = QLineEdit()
        btn = QPushButton('Proceed')
        btn.setStyleSheet('QPushButton{background-color: rgb(0, 170, 0); color: rgb(255, 255, 255)}')

        # make indigency certificate as pdf file 
        def print_putolpuno():
            doc = Document('brgy-certificate/brgy-docx-file/brgy-putolpuno.docx')

            Dictionary = {
                            "John Mark Salcedo": name.text() + ' ',
                            "100": number_of_trees.text() + ' ',
                            "niyog": tree.text() + ' ',
                            "bahay": purpose.text() + ' ',
                            "May": today.strftime('%B'), 
                            "15": today.strftime('%d'),
                            "2023": today.strftime('%Y')
                        }

            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(12)
            font.color.rgb = RGBColor(59, 56, 56)

            for i in Dictionary:
                for p in doc.paragraphs:
                    if p.text.find(i) >= 0:
                        p.text = p.text.replace(i,Dictionary[i])

            #save changed document
            
            try:
                # time.sleep(3)
                f = r"C:\Users\Angeline\Desktop\OJT-SYSTEM\calamias-system\brgy-certificate\brgy-docx-file\putolpuno-files\putolpuno_results.docx"
                file = doc.save(f)
          
                convert(f, r'C:\Users\Angeline\Desktop\OJT-SYSTEM\calamias-system\brgy-certificate\brgy-docx-file\putolpuno-files')
                print('done.....')
                df = "file:///C:/Users/Angeline/Desktop/OJT-SYSTEM/calamias-system/brgy-certificate/brgy-docx-file/putolpuno-files"
                container.dynamicCall('Navigate(const QString&)', df)

            except PermissionError:
                 QMessageBox.about(self, 'error', 'opps you need to close the file first')
               
            
        btn.clicked.connect(print_putolpuno)
        
        group_box_layout.addRow(QLabel('Pangalan: '), name)
        group_box_layout.addRow(QLabel('Piraso ng puno: '), number_of_trees)
        group_box_layout.addRow(QLabel('Uri ng Puno: '), tree)
        group_box_layout.addRow(QLabel('Saan gagamitin: '), purpose)
        group_box_layout.addRow(btn)
        box.setLayout(group_box_layout)

        WebBrowser.addWidget(box)
        WebBrowser.addWidget(container)
        main = QWidget()
        main.setLayout(WebBrowser)

        return main

    def ui7(self):
        grid = QHBoxLayout()
        label =  QtWidgets.QLabel()
        label.setMaximumSize(QtCore.QSize(700, 500))
        label.setText("")
        label.setPixmap(QtGui.QPixmap(r'img\us.jpg'))
        label.setScaledContents(True)
        label.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        label.setIndent(-1)
        label.setObjectName('logo')

        mark =  QtWidgets.QLabel()
        mark.setMaximumSize(QtCore.QSize(500, 300))
        mark.setText("")
        mark.setPixmap(QtGui.QPixmap('img/mark.png'))
        mark.setScaledContents(True)
        mark.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        mark.setIndent(-1)
        mark.setObjectName('logo')
        grid.addWidget(label)
        main = QWidget()
        main.setLayout(grid)

        return main

    def loader():
        Progress = 10
        LoadWin.progressUpdate(Progress)
        time.sleep(2)
        self.setMinimumSize(QSize(640, 480))
        self.setWindowTitle("Main Window")

        centralWidget = QWidget(self)
        self.setCentralWidget(centralWidget)
        Progress = 30
        LoadWin.progressUpdate(Progress)
        time.sleep(2)

        gridLayout = QGridLayout(self)
        centralWidget.setLayout(gridLayout)
        Progress = 50
        LoadWin.progressUpdate(Progress)
        time.sleep(2)

        title = QLabel("I am the main window", self)
        title.setAlignment(QtCore.Qt.AlignCenter)
        gridLayout.addWidget(title, 0, 0)

        Progress = 100
        LoadWin.progressUpdate(Progress)
        self.show()
        LoadWin.close()

if __name__ == '__main__':
    app = QApplication(sys.argv)
    ex = Window()
    ex.show()
    sys.exit(app.exec_())